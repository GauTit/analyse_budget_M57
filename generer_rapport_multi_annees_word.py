"""
Génération de rapport d'analyse budgétaire multi-années en format Word
à partir des données JSON enrichies

Usage:
    python generer_rapport_multi_annees_word.py <dossier_bilans>
    python generer_rapport_multi_annees_word.py "docs/bilans_multi_annees"
"""

import os
import sys
import json
import argparse
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

sys.path.insert(0, 'src')

from analysis.analyseur_multi_annees import (
    charger_bilans_multi_annees,
    comparer_bilans_annee_par_annee,
    detecter_tendances_et_anomalies,
    calculer_ratios_evolutifs
)
import pandas as pd

from generators.graphiques_evolution import generer_tous_graphiques_standard

# ============================================
# CONFIGURATION
# ============================================

DOSSIER_GRAPHIQUES = "output/graphiques_multi_annees"
FICHIER_SORTIE = "output/rapport_multi_annees.docx"


# ============================================
# CRÉATION DES STYLES WORD
# ============================================

def creer_styles_word(doc):
    """Crée les styles personnalisés pour le document Word"""

    styles = doc.styles

    # Style pour le titre principal
    if 'Titre Principal' not in [s.name for s in styles]:
        style_titre = styles.add_style('Titre Principal', WD_STYLE_TYPE.PARAGRAPH)
        style_titre.font.name = 'Arial'
        style_titre.font.size = Pt(20)
        style_titre.font.bold = True
        style_titre.font.color.rgb = RGBColor(26, 26, 26)
        style_titre.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        style_titre.paragraph_format.space_after = Pt(20)

    # Style pour les titres de section
    if 'Titre Section' not in [s.name for s in styles]:
        style_section = styles.add_style('Titre Section', WD_STYLE_TYPE.PARAGRAPH)
        style_section.font.name = 'Arial'
        style_section.font.size = Pt(16)
        style_section.font.bold = True
        style_section.font.color.rgb = RGBColor(26, 26, 26)
        style_section.paragraph_format.space_before = Pt(25)
        style_section.paragraph_format.space_after = Pt(18)

    # Style pour les sous-titres
    if 'Sous Titre' not in [s.name for s in styles]:
        style_soustitre = styles.add_style('Sous Titre', WD_STYLE_TYPE.PARAGRAPH)
        style_soustitre.font.name = 'Arial'
        style_soustitre.font.size = Pt(13)
        style_soustitre.font.bold = True
        style_soustitre.font.color.rgb = RGBColor(44, 62, 80)
        style_soustitre.paragraph_format.space_before = Pt(18)
        style_soustitre.paragraph_format.space_after = Pt(12)

    # Style pour les sous-titres dans le texte (détectés automatiquement)
    if 'Sous Titre Texte' not in [s.name for s in styles]:
        style_soustitre_texte = styles.add_style('Sous Titre Texte', WD_STYLE_TYPE.PARAGRAPH)
        style_soustitre_texte.font.name = 'Arial'
        style_soustitre_texte.font.size = Pt(11)
        style_soustitre_texte.font.bold = True
        style_soustitre_texte.font.color.rgb = RGBColor(52, 73, 94)
        style_soustitre_texte.paragraph_format.space_before = Pt(12)
        style_soustitre_texte.paragraph_format.space_after = Pt(8)

    # Style pour le corps de texte
    if 'Corps Texte' not in [s.name for s in styles]:
        style_corps = styles.add_style('Corps Texte', WD_STYLE_TYPE.PARAGRAPH)
        style_corps.font.name = 'Times New Roman'
        style_corps.font.size = Pt(10)
        style_corps.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        style_corps.paragraph_format.space_after = Pt(8)
        style_corps.paragraph_format.line_spacing = 1.25

    # Style pour les métadonnées
    if 'Metadata' not in [s.name for s in styles]:
        style_meta = styles.add_style('Metadata', WD_STYLE_TYPE.PARAGRAPH)
        style_meta.font.name = 'Times New Roman'
        style_meta.font.size = Pt(11)
        style_meta.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        style_meta.paragraph_format.space_after = Pt(5)

    # Style pour les métadonnées en gras
    if 'Metadata Gras' not in [s.name for s in styles]:
        style_meta_gras = styles.add_style('Metadata Gras', WD_STYLE_TYPE.PARAGRAPH)
        style_meta_gras.font.name = 'Times New Roman'
        style_meta_gras.font.size = Pt(12)
        style_meta_gras.font.bold = True
        style_meta_gras.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        style_meta_gras.paragraph_format.space_after = Pt(8)

    # Style pour les légendes
    if 'Legende' not in [s.name for s in styles]:
        style_legende = styles.add_style('Legende', WD_STYLE_TYPE.PARAGRAPH)
        style_legende.font.name = 'Times New Roman'
        style_legende.font.size = Pt(9)
        style_legende.font.italic = True
        style_legende.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        style_legende.paragraph_format.space_before = Pt(2)
        style_legende.paragraph_format.space_after = Pt(12)


# ============================================
# CONVERSION TEXTE EN PARAGRAPHES WORD
# ============================================

def ajouter_texte_formate(doc, texte):
    """Convertit le texte en paragraphes Word avec détection intelligente des titres"""
    if not texte:
        return

    lignes = str(texte).split('\n')

    for ligne in lignes:
        ligne = ligne.strip()
        if not ligne:
            doc.add_paragraph()
            continue

        # Nettoyer les balises HTML/Markdown
        ligne = ligne.replace('<b>', '').replace('</b>', '')
        ligne = ligne.replace('<i>', '').replace('</i>', '')
        ligne = ligne.replace('<strong>', '').replace('</strong>', '')
        ligne = ligne.replace('**', '').replace('__', '')

        # Détecter les titres de différentes façons
        est_titre = False

        # 1. Titres avec flèches (► ou ▶)
        if ligne.startswith('►') or ligne.startswith('▶'):
            ligne = ligne.replace('►', '').replace('▶', '').strip()
            est_titre = True

        # 2. Titres numérotés (ex: "1. Analyse des produits")
        elif len(ligne) > 0 and len(ligne) < 80 and ligne[0].isdigit() and '. ' in ligne[:5]:
            est_titre = True

        # 3. Titres se terminant par ":" et courts
        elif len(ligne) < 100 and ligne.endswith(':'):
            ligne = ligne.rstrip(':')
            est_titre = True

        # 4. Lignes courtes tout en majuscules (ex: "ANALYSE COMPARATIVE")
        elif len(ligne) < 80 and ligne.isupper() and not ligne.startswith('•'):
            est_titre = True

        # Appliquer le style approprié
        if est_titre:
            doc.add_paragraph(ligne, style='Sous Titre Texte')
        elif ligne.startswith('•'):
            doc.add_paragraph(ligne, style='List Bullet')
        else:
            doc.add_paragraph(ligne, style='Corps Texte')


def creer_tableau_evolution_word(doc, bilans, postes_config, par_habitant=True):
    """Crée un tableau d'évolution dans Word

    Args:
        doc: Document Word
        bilans: Liste des bilans
        postes_config: Liste de tuples (nom_poste, chemin_json)
        par_habitant: Si True, affiche les valeurs par habitant (€/hab)
    """

    annees = [str(b['annee']) for b in bilans]
    nb_colonnes = len(annees) + 2  # Poste + années + évolution

    table = doc.add_table(rows=1, cols=nb_colonnes)
    table.style = 'Light Grid Accent 1'

    # En-tête
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Poste'
    for i, annee in enumerate(annees):
        header_cells[i+1].text = annee
    header_cells[-1].text = 'Évolution'

    # Mettre en gras l'en-tête
    for cell in header_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(10)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Données
    for nom_poste, chemin in postes_config:
        row_cells = table.add_row().cells
        row_cells[0].text = nom_poste

        valeurs = []
        for i, bilan in enumerate(bilans):
            keys = chemin.split('.')
            valeur = bilan['data']

            for key in keys:
                if isinstance(valeur, dict) and key in valeur:
                    valeur = valeur[key]
                else:
                    valeur = None
                    break

            valeur_k = valeur if valeur is not None else 0

            # Convertir en €/hab si demandé
            if par_habitant:
                population = bilan['data'].get('metadata', {}).get('population')
                if population and population > 0:
                    valeur_finale = (valeur_k * 1000) / population
                    valeurs.append(valeur_finale)
                    row_cells[i+1].text = f"{valeur_finale:.0f} €/hab"
                else:
                    valeurs.append(0)
                    row_cells[i+1].text = "—"
            else:
                valeurs.append(valeur_k)
                row_cells[i+1].text = f"{valeur_k:.0f} k€"

            row_cells[i+1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # Calcul évolution
        if valeurs[0] != 0:
            evolution_pct = ((valeurs[-1] - valeurs[0]) / valeurs[0]) * 100
            row_cells[-1].text = f"{evolution_pct:+.1f}%"
        else:
            row_cells[-1].text = "—"
        row_cells[-1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # Mettre en gras la première colonne
        row_cells[0].paragraphs[0].runs[0].font.bold = True


def ajouter_encadre_points_cles(doc, titre, points):
    """Ajoute un encadré avec des points clés"""

    doc.add_paragraph(titre, style='Sous Titre')

    table = doc.add_table(rows=len(points), cols=1)
    table.style = 'Light Shading Accent 1'

    for i, point in enumerate(points):
        cell = table.rows[i].cells[0]
        p = cell.paragraphs[0]
        p.text = f"▸  {point}"
        p.style = 'Corps Texte'


# ============================================
# GÉNÉRATION DU DOCUMENT WORD
# ============================================

def generer_rapport_word_multi_annees(
    dossier_bilans: str,
    fichier_sortie: str = FICHIER_SORTIE,
    fichier_excel: str = "PROMPTS_RAPPORT_COMPLET_ENRICHIS.xlsx"
):
    """Génère le rapport Word complet d'analyse multi-années en lisant l'Excel"""

    print("\n" + "="*80)
    print("GÉNÉRATION RAPPORT D'ANALYSE MULTI-ANNÉES (WORD)")
    print("="*80 + "\n")

    # 1. Chargement des données
    print("[ÉTAPE 1/5] Chargement des données...")

    # Charger les bilans
    bilans = charger_bilans_multi_annees(dossier_bilans)
    print(f"  [OK] {len(bilans)} bilans chargés")

    # Charger l'Excel avec les analyses
    df_excel = pd.read_excel(fichier_excel)
    df_multi = df_excel[df_excel['Type_Rapport'] == 'Multi-annees'].copy()
    print(f"  [OK] Excel chargé: {len(df_multi)} analyses multi-années")

    # Créer un dictionnaire des analyses par nom de poste
    analyses_excel = {}
    for _, row in df_multi.iterrows():
        nom_poste = row['Nom_Poste']
        reponse = row.get('Reponse_Attendue', '')
        if pd.notna(reponse) and reponse.strip():
            analyses_excel[nom_poste] = reponse
        else:
            analyses_excel[nom_poste] = "Analyse non disponible"

    # Mapping entre les anciens noms et les nouveaux noms de l'Excel
    analyses = {
        'synthese_globale': analyses_excel.get('Analyse_tendances_globales', 'Analyse non disponible'),
        'poste_produits_fonctionnement': analyses_excel.get('Produits_fonctionnement_evolution', 'Analyse non disponible'),
        'poste_charges_fonctionnement': analyses_excel.get('Charges_fonctionnement_evolution', 'Analyse non disponible'),
        'poste_charges_de_personnel': analyses_excel.get('Charges_personnel_evolution', 'Analyse non disponible'),
        'poste_caf_brute': analyses_excel.get('CAF_brute_evolution', 'Analyse non disponible'),
        'poste_depenses_dequipement': analyses_excel.get('Depenses_equipement_evolution', 'Analyse non disponible'),
        'poste_encours_dette': analyses_excel.get('Encours_dette_evolution', 'Analyse non disponible')    }

    print(f"  [OK] {len([v for v in analyses.values() if v != 'Analyse non disponible'])} analyses disponibles")

    # 2. Analyse comparative
    print("\n[ÉTAPE 2/5] Analyse comparative...")
    comparaisons = comparer_bilans_annee_par_annee(bilans)
    tendances = detecter_tendances_et_anomalies(comparaisons)
    ratios = calculer_ratios_evolutifs(bilans)

    # 3. Génération des graphiques
    print("\n[ÉTAPE 3/5] Génération des graphiques...")
    os.makedirs(DOSSIER_GRAPHIQUES, exist_ok=True)
    graphiques = generer_tous_graphiques_standard(bilans, ratios['ratios_par_annee'])

    # 4. Construction du document Word
    print("\n[ÉTAPE 4/5] Construction du rapport Word...")

    os.makedirs(os.path.dirname(fichier_sortie), exist_ok=True)

    metadata = comparaisons['metadata']

    # Créer le document
    doc = Document()

    # Créer les styles personnalisés
    creer_styles_word(doc)

    # ============ PAGE DE GARDE ============
    # Titre principal
    doc.add_paragraph("RAPPORT D'ANALYSE BUDGÉTAIRE\nCOMPARATIVE MULTI-ANNÉES", style='Titre Principal')

    # Espacement
    doc.add_paragraph()
    doc.add_paragraph()

    # Métadonnées
    doc.add_paragraph(metadata['commune'], style='Metadata Gras')
    doc.add_paragraph(
        f"Période d'analyse : {metadata['periode_debut']} - {metadata['periode_fin']}",
        style='Metadata'
    )
    doc.add_paragraph(
        f"Nombre d'exercices analysés : {metadata['nb_annees']}",
        style='Metadata'
    )
    doc.add_paragraph()
    doc.add_paragraph(
        f"Rapport généré le {datetime.now().strftime('%d/%m/%Y à %H:%M')}",
        style='Metadata'
    )

    # Saut de page
    doc.add_page_break()

    # ============ SYNTHÈSE FINANCIÈRE D'ENSEMBLE ============
    doc.add_paragraph("SYNTHÈSE FINANCIÈRE D'ENSEMBLE", style='Titre Section')

    # Introduction
    intro_synthese = (
        f"Le présent rapport analyse l'évolution financière de la commune de {metadata['commune']} "
        f"sur la période {metadata['periode_debut']}-{metadata['periode_fin']}. "
        f"Cette analyse comparative porte sur {metadata['nb_annees']} exercices budgétaires et "
        f"examine les principales évolutions des sections de fonctionnement et d'investissement, "
        f"ainsi que la capacité d'autofinancement et la soutenabilité de l'endettement."
    )
    doc.add_paragraph(intro_synthese, style='Corps Texte')

    if 'synthese_globale' in analyses:
        ajouter_texte_formate(doc, analyses['synthese_globale'])

    # Points de situation à signaler
    if 'points_attention' in analyses:
        points_attention_texte = analyses.get('points_attention', '')
        points_liste = [ligne.strip().replace('•', '').strip()
                       for ligne in points_attention_texte.split('\n')
                       if ligne.strip() and ligne.strip().startswith('•')]

        if points_liste:
            ajouter_encadre_points_cles(doc, "Points de situation à signaler", points_liste)

    doc.add_page_break()

    # ============ SECTION DE FONCTIONNEMENT ============
    doc.add_paragraph("I. SECTION DE FONCTIONNEMENT", style='Titre Section')

    intro_fonct = (
        "La section de fonctionnement retrace l'ensemble des opérations courantes de la collectivité "
        "sur la période analysée. L'analyse porte sur l'évolution des produits et charges de fonctionnement."
    )
    doc.add_paragraph(intro_fonct, style='Corps Texte')

    # Évolution des produits
    doc.add_paragraph("1.1. Évolution des produits de fonctionnement", style='Sous Titre')

    # Graphique Produits vs Charges
    if 'produits_charges' in graphiques and os.path.exists(graphiques['produits_charges']):
        doc.add_picture(graphiques['produits_charges'], width=Inches(5.5))
        doc.add_paragraph(
            "Évolution comparée des produits et charges de fonctionnement (par habitant)",
            style='Legende'
        )

    # Tableau évolution
    postes_fonct = [
        ("Produits de fonctionnement", "fonctionnement.produits.total.montant_k"),
        ("Charges de fonctionnement", "fonctionnement.charges.total.montant_k"),
        ("Résultat de fonctionnement", "fonctionnement.resultat.montant_k"),
    ]
    creer_tableau_evolution_word(doc, bilans, postes_fonct)

    # Analyse des produits
    if 'poste_produits_fonctionnement' in analyses:
        ajouter_texte_formate(doc, analyses['poste_produits_fonctionnement'])

    doc.add_page_break()

    # Évolution des charges
    doc.add_paragraph("1.2. Évolution des charges de fonctionnement", style='Sous Titre')

    if 'poste_charges_fonctionnement' in analyses:
        ajouter_texte_formate(doc, analyses['poste_charges_fonctionnement'])

    if 'poste_charges_de_personnel' in analyses:
        ajouter_texte_formate(doc, analyses['poste_charges_de_personnel'])

    # Résultat de fonctionnement
    doc.add_paragraph("1.3. Résultat de fonctionnement", style='Sous Titre')

    if 'resultat' in graphiques and os.path.exists(graphiques['resultat']):
        doc.add_picture(graphiques['resultat'], width=Inches(5.5))
        doc.add_paragraph("Évolution du résultat de fonctionnement (par habitant)", style='Legende')

    transition_caf = (
        "L'évolution du résultat de fonctionnement impacte directement la capacité d'autofinancement "
        "de la collectivité."
    )
    doc.add_paragraph(transition_caf, style='Corps Texte')

    doc.add_page_break()

    # ============ CAPACITÉ D'AUTOFINANCEMENT ============
    doc.add_paragraph("II. CAPACITÉ D'AUTOFINANCEMENT", style='Titre Section')

    intro_caf = (
        "La capacité d'autofinancement constitue l'excédent dégagé par la section de fonctionnement "
        "permettant de financer les investissements et de rembourser la dette."
    )
    doc.add_paragraph(intro_caf, style='Corps Texte')

    if 'caf' in graphiques and os.path.exists(graphiques['caf']):
        doc.add_picture(graphiques['caf'], width=Inches(5.5))
        doc.add_paragraph("Évolution de la CAF brute et nette (par habitant)", style='Legende')

    postes_caf = [
        ("CAF brute", "autofinancement.caf_brute.montant_k"),
        ("CAF nette", "autofinancement.caf_nette.montant_k"),
    ]
    creer_tableau_evolution_word(doc, bilans, postes_caf)

    # Analyse CAF
    if 'poste_caf_brute' in analyses:
        ajouter_texte_formate(doc, analyses['poste_caf_brute'])

    transition_invest = (
        "La CAF nette constitue la ressource propre disponible pour financer les investissements, "
        "après remboursement du capital de la dette."
    )
    doc.add_paragraph(transition_invest, style='Corps Texte')

    doc.add_page_break()

    # ============ SECTION D'INVESTISSEMENT ============
    doc.add_paragraph("III. SECTION D'INVESTISSEMENT", style='Titre Section')

    intro_invest = (
        "La section d'investissement retrace les opérations affectant le patrimoine de la collectivité. "
        "Elle comprend les dépenses d'équipement et leurs financements."
    )
    doc.add_paragraph(intro_invest, style='Corps Texte')

    doc.add_paragraph("3.1. Dépenses d'équipement", style='Sous Titre')

    if 'depenses_equip' in graphiques and os.path.exists(graphiques['depenses_equip']):
        doc.add_picture(graphiques['depenses_equip'], width=Inches(5.5))
        doc.add_paragraph("Évolution des dépenses d'équipement (par habitant)", style='Legende')

    postes_invest = [
        ("Dépenses d'équipement", "investissement.emplois.depenses_equipement.montant_k"),
        ("Emprunts contractés", "investissement.ressources.emprunts.montant_k"),
        ("Subventions reçues", "investissement.ressources.subventions_recues.montant_k"),
    ]
    creer_tableau_evolution_word(doc, bilans, postes_invest)

    # Analyse investissement
    if 'poste_depenses_dequipement' in analyses:
        ajouter_texte_formate(doc, analyses['poste_depenses_dequipement'])

    doc.add_paragraph("3.2. Financement des investissements", style='Sous Titre')
    articulation_financement = (
        "Le financement des investissements s'articule entre autofinancement (CAF nette), "
        "emprunts contractés et subventions d'investissement reçues. L'équilibre entre ces trois sources "
        "conditionne la soutenabilité financière de la politique d'équipement."
    )
    doc.add_paragraph(articulation_financement, style='Corps Texte')

    doc.add_page_break()

    # ============ ENDETTEMENT ============
    doc.add_paragraph("IV. ENDETTEMENT", style='Titre Section')

    intro_dette = (
        "L'analyse de l'endettement porte sur l'encours de dette au 31 décembre de chaque exercice et sur "
        "la capacité de la collectivité à le rembourser dans des conditions soutenables."
    )
    doc.add_paragraph(intro_dette, style='Corps Texte')

    doc.add_paragraph("4.1. Évolution de l'encours de dette", style='Sous Titre')

    if 'dette' in graphiques and os.path.exists(graphiques['dette']):
        doc.add_picture(graphiques['dette'], width=Inches(5.5))
        doc.add_paragraph("Évolution de l'encours de dette (par habitant)", style='Legende')

    postes_dette = [
        ("Encours de dette", "endettement.encours_total.montant_k"),
    ]
    creer_tableau_evolution_word(doc, bilans, postes_dette)

    # Analyse encours
    if 'poste_encours_dette' in analyses:
        ajouter_texte_formate(doc, analyses['poste_encours_dette'])

    doc.add_paragraph("4.2. Capacité de désendettement", style='Sous Titre')

    if 'capacite_desendettement' in graphiques and os.path.exists(graphiques['capacite_desendettement']):
        doc.add_picture(graphiques['capacite_desendettement'], width=Inches(5.5))
        doc.add_paragraph(
            "Capacité de désendettement (en années) - Seuil prudentiel : 12 ans",
            style='Legende'
        )

    soutenabilite = (
        "La capacité de désendettement, ratio entre l'encours de dette et la CAF brute, "
        "constitue l'indicateur central de soutenabilité financière. Un ratio inférieur à 12 ans "
        "traduit une situation d'endettement soutenable."
    )
    doc.add_paragraph(soutenabilite, style='Corps Texte')

    doc.add_page_break()

    # ============ RATIOS FINANCIERS DE SYNTHÈSE ============
    doc.add_paragraph("V. RATIOS FINANCIERS DE SYNTHÈSE", style='Titre Section')

    intro_ratios = (
        "Les ratios financiers permettent d'apprécier la situation financière de la collectivité "
        "et sa soutenabilité à moyen terme sur la période analysée."
    )
    doc.add_paragraph(intro_ratios, style='Corps Texte')

    doc.add_paragraph("5.1. Taux d'épargne brute", style='Sous Titre')

    if 'epargne' in graphiques and os.path.exists(graphiques['epargne']):
        doc.add_picture(graphiques['epargne'], width=Inches(5.5))
        doc.add_paragraph(
            "Évolution du taux d'épargne brute (CAF brute / Produits de fonctionnement)",
            style='Legende'
        )

    explication_epargne = (
        "Le taux d'épargne brute mesure la part des produits de fonctionnement dégagée sous forme de CAF brute. "
        "Il traduit la capacité de la collectivité à dégager des marges de manœuvre financières sur son fonctionnement."
    )
    doc.add_paragraph(explication_epargne, style='Corps Texte')



    doc.add_page_break()


    # ============ SAUVEGARDE FINALE ============
    print("  [OK] Assemblage du document...")
    doc.save(fichier_sortie)

    taille_kb = os.path.getsize(fichier_sortie) / 1024

    print("\n" + "="*80)
    print(f"[OK] RAPPORT WORD GENERE AVEC SUCCES")
    print(f"  Fichier : {fichier_sortie}")
    print(f"  Taille : {taille_kb:.1f} KB")
    print(f"  Commune : {metadata['commune']}")
    print(f"  Periode : {metadata['periode_debut']}-{metadata['periode_fin']}")
    print(f"  Graphiques : {len(graphiques)}")
    print(f"  Analyses LLM : {len(analyses)}")
    print("="*80 + "\n")


def generer_rapport_word():
    """Point d'entrée pour l'import depuis workflow_complet.py"""
    dossier = "docs/bilans_multi_annees"
    fichier_sortie = "output/rapport_analyse_multi_annees.docx"
    fichier_excel = "PROMPTS_RAPPORT_COMPLET_ENRICHIS.xlsx"

    generer_rapport_word_multi_annees(dossier, fichier_sortie, fichier_excel)


# ============================================
# MAIN
# ============================================

if __name__ == "__main__":
    parser = argparse.ArgumentParser(
        description="Génère un rapport d'analyse comparative multi-années en format Word"
    )
    parser.add_argument(
        "dossier_bilans",
        help="Dossier contenant les fichiers PDF des bilans (minimum 2)"
    )
    parser.add_argument(
        "-o", "--output",
        default=FICHIER_SORTIE,
        help=f"Fichier Word de sortie (défaut: {FICHIER_SORTIE})"
    )
    parser.add_argument(
        "-e", "--excel",
        default="PROMPTS_RAPPORT_COMPLET_ENRICHIS.xlsx",
        help="Fichier Excel contenant les analyses (défaut: PROMPTS_RAPPORT_COMPLET_ENRICHIS.xlsx)"
    )

    args = parser.parse_args()

    try:
        generer_rapport_word_multi_annees(
            args.dossier_bilans,
            args.output,
            args.excel
        )
    except Exception as e:
        print(f"\n[ERREUR] {e}\n")
        import traceback
        traceback.print_exc()
        sys.exit(1)
