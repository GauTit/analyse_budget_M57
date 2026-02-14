"""
Génération de rapport d'analyse budgétaire mono-année en format Word
à partir des réponses de l'Excel et des données du JSON enrichi

Usage:
    python generer_rapport_excel_vers_word.py
"""

import os
import json
import pandas as pd
import matplotlib.pyplot as plt
import matplotlib
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from datetime import datetime

# Importer les fonctions de génération de graphiques du module PDF
from generer_rapport_excel_vers_pdf import (
    generer_graphique_repartition_produits,
    generer_graphique_repartition_charges,
    generer_graphique_comparaison_strate,
    generer_graphique_fiscalite,
    generer_graphique_ratios_financiers,
    formater_titre_poste
)

matplotlib.use('Agg')  # Backend non-interactif

# ============================================
# CONFIGURATION
# ============================================

FICHIER_EXCEL = "PROMPTS_RAPPORT_COMPLET_ENRICHIS.xlsx"
FICHIER_JSON = "output/donnees_enrichies.json"
DOSSIER_GRAPHIQUES = "output/graphiques_mono_annee"
FICHIER_SORTIE = "output/rapport_analyse_mono_annee.docx"


# ============================================
# CRÉATION DES STYLES WORD
# ============================================

def creer_styles_word(doc):
    """Crée les styles personnalisés pour le document Word"""

    styles = doc.styles

    # Style pour le titre principal
    if 'Titre Principal' not in [s.name for s in styles]:
        style_titre = styles.add_style('Titre Principal', WD_STYLE_TYPE.PARAGRAPH)
        style_titre.font.name = 'Arial'
        style_titre.font.size = Pt(20)
        style_titre.font.bold = True
        style_titre.font.color.rgb = RGBColor(26, 26, 26)
        style_titre.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        style_titre.paragraph_format.space_after = Pt(20)

    # Style pour les titres de section
    if 'Titre Section' not in [s.name for s in styles]:
        style_section = styles.add_style('Titre Section', WD_STYLE_TYPE.PARAGRAPH)
        style_section.font.name = 'Arial'
        style_section.font.size = Pt(16)
        style_section.font.bold = True
        style_section.font.color.rgb = RGBColor(26, 26, 26)
        style_section.paragraph_format.space_before = Pt(25)
        style_section.paragraph_format.space_after = Pt(18)

    # Style pour les sous-titres
    if 'Sous Titre' not in [s.name for s in styles]:
        style_soustitre = styles.add_style('Sous Titre', WD_STYLE_TYPE.PARAGRAPH)
        style_soustitre.font.name = 'Arial'
        style_soustitre.font.size = Pt(13)
        style_soustitre.font.bold = True
        style_soustitre.font.color.rgb = RGBColor(44, 62, 80)
        style_soustitre.paragraph_format.space_before = Pt(18)
        style_soustitre.paragraph_format.space_after = Pt(12)

    # Style pour les sous-titres dans le texte (détectés automatiquement)
    if 'Sous Titre Texte' not in [s.name for s in styles]:
        style_soustitre_texte = styles.add_style('Sous Titre Texte', WD_STYLE_TYPE.PARAGRAPH)
        style_soustitre_texte.font.name = 'Arial'
        style_soustitre_texte.font.size = Pt(11)
        style_soustitre_texte.font.bold = True
        style_soustitre_texte.font.color.rgb = RGBColor(52, 73, 94)
        style_soustitre_texte.paragraph_format.space_before = Pt(12)
        style_soustitre_texte.paragraph_format.space_after = Pt(8)

    # Style pour le corps de texte
    if 'Corps Texte' not in [s.name for s in styles]:
        style_corps = styles.add_style('Corps Texte', WD_STYLE_TYPE.PARAGRAPH)
        style_corps.font.name = 'Times New Roman'
        style_corps.font.size = Pt(10)
        style_corps.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        style_corps.paragraph_format.space_after = Pt(8)
        style_corps.paragraph_format.line_spacing = 1.25

    # Style pour les métadonnées
    if 'Metadata' not in [s.name for s in styles]:
        style_meta = styles.add_style('Metadata', WD_STYLE_TYPE.PARAGRAPH)
        style_meta.font.name = 'Times New Roman'
        style_meta.font.size = Pt(11)
        style_meta.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        style_meta.paragraph_format.space_after = Pt(5)

    # Style pour les métadonnées en gras
    if 'Metadata Gras' not in [s.name for s in styles]:
        style_meta_gras = styles.add_style('Metadata Gras', WD_STYLE_TYPE.PARAGRAPH)
        style_meta_gras.font.name = 'Times New Roman'
        style_meta_gras.font.size = Pt(12)
        style_meta_gras.font.bold = True
        style_meta_gras.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        style_meta_gras.paragraph_format.space_after = Pt(8)

    # Style pour les légendes
    if 'Legende' not in [s.name for s in styles]:
        style_legende = styles.add_style('Legende', WD_STYLE_TYPE.PARAGRAPH)
        style_legende.font.name = 'Times New Roman'
        style_legende.font.size = Pt(9)
        style_legende.font.italic = True
        style_legende.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        style_legende.paragraph_format.space_before = Pt(2)
        style_legende.paragraph_format.space_after = Pt(12)


# ============================================
# GÉNÉRATION DES GRAPHIQUES
# ============================================

def generer_tous_graphiques(data_json):
    """Génère tous les graphiques à partir du JSON"""

    os.makedirs(DOSSIER_GRAPHIQUES, exist_ok=True)

    graphiques = {}

    print("\n[GRAPHIQUES] Génération en cours...")

    # Graphique 1: Répartition des produits
    fichier = os.path.join(DOSSIER_GRAPHIQUES, "repartition_produits.png")
    generer_graphique_repartition_produits(data_json, fichier)
    graphiques['repartition_produits'] = fichier
    print(f"  [OK] Repartition des produits")

    # Graphique 2: Répartition des charges
    fichier = os.path.join(DOSSIER_GRAPHIQUES, "repartition_charges.png")
    generer_graphique_repartition_charges(data_json, fichier)
    graphiques['repartition_charges'] = fichier
    print(f"  [OK] Repartition des charges")

    # Graphique 3: Comparaison strate
    fichier = os.path.join(DOSSIER_GRAPHIQUES, "comparaison_strate.png")
    generer_graphique_comparaison_strate(data_json, fichier)
    graphiques['comparaison_strate'] = fichier
    print(f"  [OK] Comparaison avec strate")

    # Graphique 4: Fiscalité
    fichier = os.path.join(DOSSIER_GRAPHIQUES, "fiscalite.png")
    result = generer_graphique_fiscalite(data_json, fichier)
    if result:
        graphiques['fiscalite'] = fichier
        print(f"  [OK] Comparaison fiscalite")

    # Graphique 5: Ratios financiers
    fichier = os.path.join(DOSSIER_GRAPHIQUES, "ratios_financiers.png")
    result = generer_graphique_ratios_financiers(data_json, fichier)
    if result:
        graphiques['ratios_financiers'] = fichier
        print(f"  [OK] Ratios financiers")

    print(f"\n  Total: {len(graphiques)} graphiques générés\n")

    return graphiques


# ============================================
# CONVERSION TEXTE EN PARAGRAPHES WORD
# ============================================

def ajouter_texte_formate(doc, texte):
    """Convertit le texte en paragraphes Word avec détection intelligente des titres"""
    if not texte or pd.isna(texte):
        return

    lignes = str(texte).split('\n')

    for ligne in lignes:
        ligne = ligne.strip()
        if not ligne:
            doc.add_paragraph()  # Ligne vide
            continue

        # Nettoyer les balises HTML/Markdown
        ligne = ligne.replace('<b>', '').replace('</b>', '')
        ligne = ligne.replace('<i>', '').replace('</i>', '')
        ligne = ligne.replace('<strong>', '').replace('</strong>', '')
        ligne = ligne.replace('**', '').replace('__', '')

        # Détecter les titres de différentes façons
        est_titre = False

        # 1. Titres avec flèches (► ou ▶)
        if ligne.startswith('►') or ligne.startswith('▶'):
            ligne = ligne.replace('►', '').replace('▶', '').strip()
            est_titre = True

        # 2. Titres numérotés (ex: "1. Analyse des produits")
        elif len(ligne) > 0 and len(ligne) < 80 and ligne[0].isdigit() and '. ' in ligne[:5]:
            est_titre = True

        # 3. Titres se terminant par ":" et courts
        elif len(ligne) < 100 and ligne.endswith(':'):
            ligne = ligne.rstrip(':')
            est_titre = True

        # 4. Lignes courtes tout en majuscules (ex: "ANALYSE COMPARATIVE")
        elif len(ligne) < 80 and ligne.isupper() and not ligne.startswith('•'):
            est_titre = True

        # Appliquer le style approprié
        if est_titre:
            doc.add_paragraph(ligne, style='Sous Titre Texte')
        elif ligne.startswith('•'):
            doc.add_paragraph(ligne, style='List Bullet')
        else:
            doc.add_paragraph(ligne, style='Corps Texte')


# ============================================
# GÉNÉRATION DU DOCUMENT WORD
# ============================================

def generer_rapport_word():
    """Génère le rapport Word complet"""

    print("\n" + "="*80)
    print("GÉNÉRATION RAPPORT D'ANALYSE BUDGÉTAIRE MONO-ANNÉE (WORD)")
    print("="*80 + "\n")

    # 1. Chargement des données
    print("[ÉTAPE 1/4] Chargement des données...")

    # Charger l'Excel
    df_excel = pd.read_excel(FICHIER_EXCEL)
    df_mono = df_excel[df_excel['Type_Rapport'] == 'Mono-annee'].copy()
    print(f"  [OK] Excel charge: {len(df_mono)} analyses mono-annee")

    # Charger le JSON
    with open(FICHIER_JSON, 'r', encoding='utf-8') as f:
        data_json = json.load(f)
    print(f"  [OK] JSON charge: {data_json['metadata']['commune']} - {data_json['metadata']['exercice']}")

    # 2. Génération des graphiques
    print("\n[ÉTAPE 2/4] Génération des graphiques...")
    graphiques = generer_tous_graphiques(data_json)

    # 3. Construction du document Word
    print("[ÉTAPE 3/4] Construction du rapport Word...")

    os.makedirs(os.path.dirname(FICHIER_SORTIE), exist_ok=True)

    metadata = data_json['metadata']

    # Créer le document
    doc = Document()

    # Créer les styles personnalisés
    creer_styles_word(doc)

    # ============ PAGE DE GARDE ============
    # Titre principal
    titre = doc.add_paragraph("RAPPORT D'ANALYSE BUDGÉTAIRE\nMONO-ANNÉE", style='Titre Principal')

    # Espacement
    doc.add_paragraph()
    doc.add_paragraph()

    # Métadonnées
    doc.add_paragraph(metadata['commune'], style='Metadata Gras')
    doc.add_paragraph(f"Exercice {metadata['exercice']}", style='Metadata')
    doc.add_paragraph(
        f"Population : {metadata['population']} habitants",
        style='Metadata'
    )
    doc.add_paragraph(
        f"Strate démographique : {metadata['strate']['libelle']}",
        style='Metadata'
    )
    doc.add_paragraph()
    doc.add_paragraph(
        f"Rapport généré le {datetime.now().strftime('%d/%m/%Y à %H:%M')}",
        style='Metadata'
    )

    # Saut de page
    doc.add_page_break()

    # ============ PRÉSENTATION DE LA COLLECTIVITÉ ============
    doc.add_paragraph("PRÉSENTATION DE LA COLLECTIVITÉ", style='Titre Section')

    # Informations générales
    doc.add_paragraph("Identification", style='Sous Titre')

    texte_identification = (
        f"La commune de {metadata['commune']} compte {metadata['population']} habitants "
        f"au titre de l'exercice {metadata['exercice']}. "
        f"Elle se situe dans la strate démographique des communes de {metadata['strate']['libelle']}."
    )
    doc.add_paragraph(texte_identification, style='Corps Texte')

    # Contexte de l'analyse
    doc.add_paragraph("Objet de l'analyse", style='Sous Titre')

    texte_objet = (
        "Le présent rapport d'analyse budgétaire a pour objet d'apprécier la situation financière "
        "de la collectivité au titre de l'exercice budgétaire considéré. L'analyse porte sur les "
        "opérations de fonctionnement et d'investissement, la capacité d'autofinancement, l'endettement "
        "et les équilibres financiers. Les données sont systématiquement comparées à la moyenne "
        "de la strate démographique de référence."
    )
    doc.add_paragraph(texte_objet, style='Corps Texte')

    # Méthodologie
    doc.add_paragraph("Méthodologie", style='Sous Titre')

    texte_methodo = (
        "L'analyse financière s'appuie sur les données des budgets exécutés par les communes "
        "dont la source provient de la Direction Générale des Finances Publiques (DGFiP)."
        " Elle respecte la nomenclature comptable M57. Les comparaisons avec la strate démographique permettent "
        "de situer la collectivité par rapport aux communes de taille comparable. Les ratios de niveau "
        "sont exprimés en euros par habitant. Les ratios de structure sont exprimés en %."
    )
    doc.add_paragraph(texte_methodo, style='Corps Texte')

    # Saut de page
    doc.add_page_break()

    # ============ SYNTHÈSE GLOBALE ============
    synthese = df_mono[df_mono['Section'] == 'Synthese_globale']
    if not synthese.empty:
        doc.add_paragraph("SYNTHÈSE FINANCIÈRE GLOBALE", style='Titre Section')

        for _, row in synthese.iterrows():
            # Ajouter le titre du poste budgétaire
            doc.add_paragraph(formater_titre_poste(row['Nom_Poste']), style='Sous Titre Texte')
            ajouter_texte_formate(doc, row['Reponse_Attendue'])

        doc.add_page_break()

    # ============ SECTION DE FONCTIONNEMENT ============
    doc.add_paragraph("I. SECTION DE FONCTIONNEMENT", style='Titre Section')

    # Introduction générale au fonctionnement
    intro_fonct = (
        "La section de fonctionnement retrace l'ensemble des opérations courantes de la collectivité. "
        "Elle se caractérise par les produits (recettes) et les charges (dépenses) nécessaires au "
        "fonctionnement des services publics locaux."
    )
    doc.add_paragraph(intro_fonct, style='Corps Texte')

    # 1.1 Produits de fonctionnement
    doc.add_paragraph("1.1. Produits de fonctionnement", style='Sous Titre')

    # Analyses produits
    fonctionnement = df_mono[df_mono['Section'] == 'Fonctionnement']
    produits_analyses = fonctionnement[fonctionnement['Nom_Poste'].str.contains('produit|Produit', case=False, na=False)]

    for _, row in produits_analyses.iterrows():
        # Ajouter le titre du poste budgétaire
        doc.add_paragraph(formater_titre_poste(row['Nom_Poste']), style='Sous Titre Texte')
        ajouter_texte_formate(doc, row['Reponse_Attendue'])

    # Graphique répartition produits
    if 'repartition_produits' in graphiques:
        doc.add_picture(graphiques['repartition_produits'], width=Inches(5.5))
        doc.add_paragraph("Graphique 1 – Répartition des recettes réelles de fonctionnement", style='Legende')

    doc.add_page_break()

    # 1.2 Charges de fonctionnement
    doc.add_paragraph("1.2. Charges de fonctionnement", style='Sous Titre')

    # Analyses charges
    charges_analyses = fonctionnement[fonctionnement['Nom_Poste'].str.contains('charge|Charge', case=False, na=False)]

    for _, row in charges_analyses.iterrows():
        # Ajouter le titre du poste budgétaire
        doc.add_paragraph(formater_titre_poste(row['Nom_Poste']), style='Sous Titre Texte')
        ajouter_texte_formate(doc, row['Reponse_Attendue'])

    # Graphique répartition charges
    if 'repartition_charges' in graphiques:
        doc.add_picture(graphiques['repartition_charges'], width=Inches(5.5))
        doc.add_paragraph("Graphique 2 – Répartition des dépenses réelles de fonctionnement", style='Legende')

    doc.add_page_break()

    # 1.3 Analyse comparative
    doc.add_paragraph("1.3. Analyse comparative", style='Sous Titre')

    # Texte d'analyse comparative
    texte_comparatif = (
        "Le positionnement de la commune par rapport à la moyenne de sa strate démographique "
        "permet d'apprécier le niveau relatif de ses produits et charges de fonctionnement. "
        "Cette comparaison constitue un élément d'appréciation de la structure financière communale."
    )
    doc.add_paragraph(texte_comparatif, style='Corps Texte')

    # Graphique comparaison strate
    if 'comparaison_strate' in graphiques:
        doc.add_picture(graphiques['comparaison_strate'], width=Inches(5.5))
        doc.add_paragraph(
            "Graphique 3 – Comparaison avec la moyenne de strate",
            style='Legende'
        )

    # Graphique fiscalité
    if 'fiscalite' in graphiques:
        texte_fiscalite = (
            "L'analyse de la pression fiscale locale s'effectue par comparaison des taux communaux "
            "avec les taux moyens constatés dans la strate."
        )
        doc.add_paragraph(texte_fiscalite, style='Corps Texte')

        doc.add_picture(graphiques['fiscalite'], width=Inches(5.5))
        doc.add_paragraph(
            "Graphique 4 – Comparaison des taux de fiscalité locale",
            style='Legende'
        )

    doc.add_page_break()

    # ============ AUTOFINANCEMENT ============
    autofinancement = df_mono[df_mono['Section'] == 'Autofinancement']
    if not autofinancement.empty:
        doc.add_paragraph("II. CAPACITÉ D'AUTOFINANCEMENT", style='Titre Section')

        intro_caf = (
            "La capacité d'autofinancement constitue l'excédent dégagé par la section de fonctionnement "
            "permettant de financer les investissements et de rembourser la dette."
        )
        doc.add_paragraph(intro_caf, style='Corps Texte')

        for _, row in autofinancement.iterrows():
            # Ajouter le titre du poste budgétaire
            doc.add_paragraph(formater_titre_poste(row['Nom_Poste']), style='Sous Titre Texte')
            ajouter_texte_formate(doc, row['Reponse_Attendue'])

        doc.add_page_break()

    # ============ INVESTISSEMENT ============
    investissement = df_mono[df_mono['Section'] == 'Investissement']
    if not investissement.empty:
        doc.add_paragraph("III. SECTION D'INVESTISSEMENT", style='Titre Section')

        intro_invest = (
            "La section d'investissement retrace les opérations affectant le patrimoine de la collectivité. "
            "Elle comprend les dépenses d'équipement et leurs financements."
        )
        doc.add_paragraph(intro_invest, style='Corps Texte')

        for _, row in investissement.iterrows():
            # Ajouter le titre du poste budgétaire
            doc.add_paragraph(formater_titre_poste(row['Nom_Poste']), style='Sous Titre Texte')
            ajouter_texte_formate(doc, row['Reponse_Attendue'])

        doc.add_page_break()

    # ============ ENDETTEMENT ============
    endettement = df_mono[df_mono['Section'] == 'Endettement']
    if not endettement.empty:
        doc.add_paragraph("IV. ENDETTEMENT", style='Titre Section')

        intro_dette = (
            "L'analyse de l'endettement porte sur l'encours de dette au 31 décembre et sur "
            "la capacité de la collectivité à le rembourser dans des conditions soutenables."
        )
        doc.add_paragraph(intro_dette, style='Corps Texte')

        for _, row in endettement.iterrows():
            # Ajouter le titre du poste budgétaire
            doc.add_paragraph(formater_titre_poste(row['Nom_Poste']), style='Sous Titre Texte')
            ajouter_texte_formate(doc, row['Reponse_Attendue'])

        doc.add_page_break()



    # ============ SAUVEGARDE FINALE ============
    print("  [OK] Assemblage du document...")
    doc.save(FICHIER_SORTIE)

    # 4. Statistiques finales
    print("\n[ETAPE 4/4] Generation terminee")

    taille_kb = os.path.getsize(FICHIER_SORTIE) / 1024

    print("\n" + "="*80)
    print(f"[OK] RAPPORT WORD GENERE AVEC SUCCES")
    print(f"  Fichier : {FICHIER_SORTIE}")
    print(f"  Taille : {taille_kb:.1f} KB")
    print(f"  Commune : {metadata['commune']}")
    print(f"  Exercice : {metadata['exercice']}")
    print(f"  Analyses : {len(df_mono)}")
    print(f"  Graphiques : {len(graphiques)}")
    print("="*80 + "\n")


# ============================================
# MAIN
# ============================================

if __name__ == "__main__":
    try:
        generer_rapport_word()
    except Exception as e:
        print(f"\n[ERREUR]: {e}\n")
        import traceback
        traceback.print_exc()
